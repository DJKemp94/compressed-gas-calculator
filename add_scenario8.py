import docx

doc = docx.Document("Worked Examples.docx")

# Prepare new paragraphs for Question
q_title = "Scenario 8: Mixed Gas Cylinders & Multiple Independent Risks"
q_text_1 = "You are using a standard BOC size 'K' (50L) cylinder of a mixed gas comprising 10% Carbon Dioxide, 10% Hydrogen, and 80% Nitrogen (Balance). The initial pressure is 200 bar. The cylinder is located in a small 30 m³ instrument room."
q_bullets = [
    "When assessing mixed gases in the calculator, you must assess each component gas independently. First, what is the total volume of gas in the cylinder?",
    "Next, calculate the total Oxygen depletion caused by the entire volume of gas displacing the room air. Is Oxygen detection required?",
    "Now, calculate the individual gas volumes for the CO₂ (10% of total) and Hydrogen (10% of total), and determine their specific room concentrations.",
    "Based on these individual concentrations, what specific detectors does this single cylinder require, and where should they be placed?"
]

# Prepare new paragraphs for Answer
a_title = "Scenario 8: Mixed Gas Cylinders & Multiple Independent Risks"
a_bullets = [
    "Total Gas Volume:\t9.87 m³. (Calculated using 200 bar × 50L).",
    "Oxygen Depletion:\tYes, O₂ detection is absolutely required. The total gas release (9.87 m³) displaces air, resulting in a lethal room O₂ concentration of 14.06%.",
    "CO₂ Hazard:\tThe CO₂ gas volume is 0.987 m³ (10% of total). This results in a room CO₂ concentration of 3.29%. This is extremely hazardous (approaching the 4% IDLH threshold) and requires CO₂ detection.",
    "Hydrogen Hazard:\tThe H₂ gas volume is 0.987 m³ (10% of total). This results in a room H₂ concentration of 3.29%. The Lower Explosive Limit (LEL) for Hydrogen is 4%, and detection is required at 25% of the LEL (1%). Therefore, Flammable detection is required.",
    "Detection Placements:\tThis single cylinder creates three distinct hazards requiring three different sensors: an O₂ sensor at breathing height, a CO₂ sensor at a low level (heavy gas), and a flammable (H₂) sensor at the ceiling (very light gas)."
]

# Find where to insert in Part 1 (Right before Part 2)
part2_idx = -1
part3_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Part 2: Answers" in p.text:
        part2_idx = i
    elif "Part 3: Formulas Used" in p.text:
        part3_idx = i

# Insert answered scenario before Part 3
if part3_idx != -1:
    p_ans_title = doc.paragraphs[part3_idx].insert_paragraph_before(a_title)
    p_ans_title.runs[0].bold = True
    for bullet in a_bullets:
        doc.paragraphs[part3_idx].insert_paragraph_before(bullet, style='List Paragraph')

# Insert questioned scenario before Part 2
if part2_idx != -1:
    # Need to find the exact paragraph of Part 2 to insert before, it shifts when we insert above if we are not careful
    # But since part2_idx is before part3_idx, if we inserted at part3_idx, part2_idx is unchanged.
    p_q_title = doc.paragraphs[part2_idx].insert_paragraph_before(q_title)
    p_q_title.runs[0].bold = True
    p_q_text = doc.paragraphs[part2_idx].insert_paragraph_before(q_text_1)
    
    for bullet in q_bullets:
        doc.paragraphs[part2_idx].insert_paragraph_before(bullet, style='List Paragraph')
        
    # Add a blank line for spacing
    doc.paragraphs[part2_idx].insert_paragraph_before("")

doc.save("Worked Examples.docx")
print("Successfully added Scenario 8 (Mixed Gases) to Worked Examples.docx")
